import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
import tempfile
import asyncio
import edge_tts
from docx import Document
import io
import os
import base64
from typing import List, Dict

# Optional OCR / clipboard imports
try:
    import pytesseract
    from PIL import Image, ImageGrab
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# ─────────────────────────────────────
# MastersHub2023 branding & colours
# ─────────────────────────────────────
MH_PRIMARY = "#004AAD"   # MastersHub blue
MH_ACCENT = "#CEFA4A"    # MastersHub neon green
MH_BG_SOFT = "#0B1020"   # dark-ish background accent

st.set_page_config(
    page_title="MastersHub2023 — Doc Reader with HD Voice",
    layout="wide",
)

st.markdown(
    f"""
    <div style="background:{MH_BG_SOFT};padding:12px 18px;border-radius:14px;
                display:flex;justify-content:space-between;align-items:center;">
        <div>
            <h1 style="margin:0;color:{MH_ACCENT};font-size:1.7rem;">
                🗂️ MastersHub2023 — Document Reader with HD Voice
            </h1>
            <p style="margin:2px 0 0 0;color:#f5f5f5;font-size:0.9rem;">
                Upload DOCX, PDF, CSV, or XLSX. View content & tables.
                Read aloud with natural neural voices.
            </p>
        </div>
        <span style="color:#ddd;font-size:0.8rem;">v2.1</span>
    </div>
    """,
    unsafe_allow_html=True,
)

# ─────────────────────────────────────
# Helpers: document reading
# ─────────────────────────────────────
def read_docx(file_bytes):
    doc = Document(file_bytes)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        rows = []
        max_cols = max(len(r.cells) for r in t.rows) if t.rows else 0
        for r in t.rows:
            rows.append([c.text for c in r.cells] + [""] * (max_cols - len(r.cells)))
        df = pd.DataFrame(rows)
        if len(df) > 1:
            df.columns = df.iloc[0].fillna("").astype(str)
            df = df[1:].reset_index(drop=True)
        tables.append(df)
    return "\n\n".join(paragraphs), tables


def read_pdf(file_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tf:
        tf.write(file_bytes.read())
        temp_path = tf.name

    doc = fitz.open(temp_path)
    pages = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=180)
        img_bytes = pix.tobytes("png")
        text = page.get_text("text")
        pages.append({"index": i + 1, "image": img_bytes, "text": text})
    doc.close()

    try:
        os.remove(temp_path)
    except Exception:
        pass

    return pages


def read_tabular(file_bytes, filetype="csv"):
    if filetype == "csv":
        df = pd.read_csv(file_bytes)
        return [df]
    elif filetype == "xlsx":
        xls = pd.ExcelFile(file_bytes)
        tables = []
        for sheet in xls.sheet_names:
            tables.append(xls.parse(sheet_name=sheet))
        return tables
    else:
        return []

# ─────────────────────────────────────
# Voice utilities (edge-tts)
# ─────────────────────────────────────
async def fetch_voices_async():
    """Get full list of Microsoft voices."""
    voices = await edge_tts.list_voices()
    return voices


def load_voice_catalog() -> List[Dict]:
    """Load & normalize voice list, store in session_state."""
    if "voice_catalog" in st.session_state:
        return st.session_state["voice_catalog"]

    try:
        raw = asyncio.run(fetch_voices_async())
        catalog = []
        for v in raw:
            catalog.append(
                {
                    "short_name": v.get("ShortName") or v.get("Name") or "",
                    "locale": v.get("Locale") or "",
                    "gender": v.get("Gender") or "",
                    "style_list": v.get("StyleList") or [],
                }
            )
        st.session_state["voice_catalog"] = catalog
        return catalog
    except Exception as e:
        st.warning(
            f"Could not fetch full voice list automatically ({e}). "
            "Using a small fallback list instead."
        )
        catalog = [
            {"short_name": n, "locale": "en-US", "gender": "Female", "style_list": []}
            for n in [
                "en-US-AriaNeural",
                "en-US-JennyNeural",
                "en-US-GuyNeural",
                "en-GB-LibbyNeural",
                "en-GB-RyanNeural",
                "en-AU-NatashaNeural",
                "en-IN-NeerjaNeural",
            ]
        ]
        st.session_state["voice_catalog"] = catalog
        return catalog


async def synthesize_to_file(text, voice, rate, volume, out_suffix=".mp3"):
    communicate = edge_tts.Communicate(text, voice=voice, rate=rate, volume=volume)
    with tempfile.NamedTemporaryFile(delete=False, suffix=out_suffix) as tf:
        out_path = tf.name
    await communicate.save(out_path)
    return out_path


async def stream_chunk_to_file(
    text, voice, rate, volume, fhandle, progress=None, cur=0, total=1
):
    """Stream a chunk to an open file handle, update progress."""
    communicate = edge_tts.Communicate(text, voice=voice, rate=rate, volume=volume)
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            fhandle.write(chunk["data"])
    if progress:
        progress.progress(min(1.0, (cur + 1) / total))


def chunk_text(text: str, max_len: int = 600):
    """Split long text into approximate sentence chunks."""
    parts = []
    buf = []
    length = 0
    for token in text.split(" "):
        if length + len(token) + 1 > max_len:
            parts.append(" ".join(buf))
            buf = [token]
            length = len(token) + 1
        else:
            buf.append(token)
            length += len(token) + 1
    if buf:
        parts.append(" ".join(buf))
    return [p.strip() for p in parts if p.strip()]


# ─────────────────────────────────────
# Sidebar: Voice selection (MastersHub style)
# ─────────────────────────────────────
st.sidebar.subheader("🔊 Voice Settings — MastersHub2023")

voice_catalog = load_voice_catalog()

# rate / volume controls
rate = st.sidebar.slider("Rate (−50% to +50%)", -50, 50, 0, format="%d%%")
volume = st.sidebar.slider("Volume (−50% to +50%)", -50, 50, 0, format="%d%%")
rate_str = f"{rate:+d}%"
volume_str = f"{volume:+d}%"


with st.sidebar.expander("Choose a Voice (HD style list)", expanded=False):
    q = st.text_input("Search (name / locale / style):", "")
    locales = sorted({v["locale"] for v in voice_catalog if v.get("locale")})
    locale_filter = st.selectbox("Locale", ["(all)"] + locales, index=0)
    gender_filter = st.selectbox("Gender", ["(all)", "Female", "Male"], index=0)

    filtered = []
    for v in voice_catalog:
        if locale_filter != "(all)" and v["locale"] != locale_filter:
            continue
        if gender_filter != "(all)" and v["gender"].lower() != gender_filter.lower():
            continue

        line = f"{v['short_name']} {v['locale']} {' '.join(v['style_list'])}"
        if q.strip().lower() not in line.lower():
            continue
        filtered.append(v)

    if not filtered:
        st.info("No voices match these filters.")
    else:
        shown = filtered[:40]  # avoid huge list

        for i, v in enumerate(shown, start=1):
            name = v["short_name"]
            locale = v["locale"]
            gender = v["gender"] or ""
            styles = ", ".join(v["style_list"]) if v["style_list"] else "HD Voice"

            c1, c2, c3 = st.columns([0.15, 0.55, 0.30])
            with c1:
                st.markdown(
                    f"""
                    <div style="width:38px;height:38px;border-radius:50%;
                                background:{MH_PRIMARY};display:flex;
                                align-items:center;justify-content:center;
                                color:white;font-size:0.9rem;">
                        🗣️
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c2:
                st.markdown(
                    f"""
                    <div style="line-height:1.1;">
                        <span style="font-weight:600;">{name}</span>
                        <span style="opacity:0.8;"> ({locale})</span>
                        <span style="color:{MH_ACCENT};font-size:0.7rem;margin-left:6px;">
                            New
                        </span><br/>
                        <span style="font-size:0.8rem;opacity:0.9;">
                            {styles}
                        </span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c3:
                if st.button(f"▶ Preview", key=f"pv_{name}_{i}"):
                    st.session_state["preview_voice"] = name
                if st.button(f"Use", key=f"use_{name}_{i}"):
                    st.session_state["selected_voice"] = name

        st.caption(f"Showing {len(shown)} of {len(filtered)} match(es).")

# selected voice (fallback)
selected_voice = st.session_state.get("selected_voice", "en-US-JennyNeural")
st.sidebar.success(f"Selected voice: {selected_voice}")

# voice preview player (auto-play)
if "preview_voice" in st.session_state:
    pv = st.session_state.pop("preview_voice")
    with st.spinner(f"Synthesizing preview for {pv}…"):
        sample_text = (
            f"Hello from {pv}. "
            "This is a short preview from MastersHub twenty twenty three."
        )
        out_path = asyncio.run(
            synthesize_to_file(sample_text, pv, rate_str, volume_str, ".mp3")
        )
        audio_bytes = open(out_path, "rb").read()

    audio_b64 = base64.b64encode(audio_bytes).decode("utf-8")
    audio_html = f"""
    <audio controls autoplay style="width: 100%;">
        <source src="data:audio/mp3;base64,{audio_b64}" type="audio/mpeg">
        Your browser does not support the audio element.
    </audio>
    """
    st.sidebar.markdown(audio_html, unsafe_allow_html=True)

# ─────────────────────────────────────
# File uploader (documents)
# ─────────────────────────────────────
uploaded = st.file_uploader(
    "Upload a document",
    type=["docx", "pdf", "csv", "xlsx"],
)
text_to_read = ""

if uploaded:
    suffix = uploaded.name.lower().split(".")[-1]

    if suffix == "docx":
        st.subheader(f"📄 {uploaded.name}")
        text, tables = read_docx(uploaded)
        if text:
            st.markdown("### Extracted Text")
            st.write(text)
            text_to_read = text[:12000]

        if tables:
            st.markdown("### Tables (original structure preserved)")
            for i, df in enumerate(tables, start=1):
                st.markdown(f"**Table {i}**")
                st.dataframe(df, use_container_width=True)

    elif suffix == "pdf":
        st.subheader(f"📕 {uploaded.name}")
        pages = read_pdf(uploaded)
        st.info(
            "Showing each page as an image to preserve exact table layout. "
            "Plain text is also provided where available."
        )
        for p in pages:
            st.markdown(f"#### Page {p['index']}")
            st.image(p["image"], use_column_width=True)
            if p["text"].strip():
                with st.expander("Show extracted plain text"):
                    st.text(p["text"])
                    if not text_to_read:
                        text_to_read = p["text"][:12000]

    elif suffix in ("csv", "xlsx"):
        st.subheader(f"📊 {uploaded.name}")
        if suffix == "csv":
            tables = read_tabular(uploaded, "csv")
        else:
            tables = read_tabular(uploaded, "xlsx")

        for i, df in enumerate(tables, start=1):
            st.markdown(f"**Sheet/Table {i}**")
            st.dataframe(df, use_container_width=True)

        if tables:
            buf = io.StringIO()
            tables[0].to_csv(buf, index=False)
            text_to_read = buf.getvalue()[:12000]

# ─────────────────────────────────────
# Image reader (upload + clipboard + OCR + TTS)
# ─────────────────────────────────────
st.markdown("---")
st.header("🖼️ Image Reader (beta)")

if "clipboard_images" not in st.session_state:
    st.session_state["clipboard_images"] = []  # store PNG bytes

col_up, col_clip = st.columns([0.6, 0.4])

with col_up:
    images = st.file_uploader(
        "Upload up to 10 images (PNG / JPG / JPEG)",
        type=["png", "jpg", "jpeg"],
        accept_multiple_files=True,
    )

with col_clip:
    if st.button("📋 Paste image from clipboard"):
        if not OCR_AVAILABLE:
            st.warning(
                "Clipboard image support needs Pillow (PIL). "
                "Install with: pip install pillow"
            )
        else:
            try:
                clip = ImageGrab.grabclipboard()
                if isinstance(clip, Image.Image):
                    buf = io.BytesIO()
                    clip.save(buf, format="PNG")
                    buf.seek(0)
                    st.session_state["clipboard_images"].append(buf.getvalue())
                    st.success("Image pasted from clipboard.")
                elif isinstance(clip, list) and clip:
                    added = 0
                    for item in clip:
                        if isinstance(item, str) and os.path.isfile(item):
                            try:
                                img = Image.open(item)
                                buf = io.BytesIO()
                                img.save(buf, format="PNG")
                                buf.seek(0)
                                st.session_state["clipboard_images"].append(buf.getvalue())
                                added += 1
                            except Exception:
                                pass
                    if added:
                        st.success(f"{added} image(s) loaded from clipboard file paths.")
                    else:
                        st.warning("Clipboard does not contain an image.")
                else:
                    st.warning("Clipboard does not contain an image.")
            except Exception as e:
                st.warning(f"Could not access clipboard: {e}")

st.caption(
    "Tip: Take a screenshot (Win + Shift + S), then click '📋 Paste image from clipboard' "
    "to pull it directly into MastersHub2023."
)

# Combine uploaded + clipboard images into one ordered list
image_sources = []  # (label, kind, data)

if images:
    if len(images) > 10:
        st.info("You uploaded more than 10 images. Only the first 10 will be used.")
        images = images[:10]
    for uf in images:
        image_sources.append((f"Uploaded: {uf.name}", "uploaded", uf))

for i, b in enumerate(st.session_state["clipboard_images"], start=1):
    image_sources.append((f"Clipboard image {i}", "clipboard", b))

image_texts = []

if image_sources:
    if not OCR_AVAILABLE:
        st.warning(
            "OCR is not available. Install Tesseract OCR and the Python packages "
            "`pytesseract` and `pillow` to enable text extraction."
        )

    for idx, (label, kind, data) in enumerate(image_sources, start=1):
        st.markdown(f"**Image {idx} — {label}**")

        if kind == "uploaded":
            st.image(data, use_column_width=True)
            pil_img = Image.open(data) if OCR_AVAILABLE else None
        else:  # clipboard bytes
            buf = io.BytesIO(data)
            pil_img = Image.open(buf) if OCR_AVAILABLE else None
            st.image(pil_img, use_column_width=True)

        if OCR_AVAILABLE and pil_img is not None:
            text = pytesseract.image_to_string(pil_img).strip()
            if text:
                st.text_area(
                    f"Extracted text for image {idx}",
                    value=text,
                    height=90,
                )
                image_texts.append(f"Image {idx}: {text}")
            else:
                st.caption("No text detected in this image.")

if image_sources and image_texts:
    if st.button("▶️ Read Text from Images in Order"):
        combined = "\n\n".join(image_texts)
        chunks = chunk_text(combined, max_len=700)
        img_progress = st.progress(0.0, text="Synthesizing from images…")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tf:
            img_out_path = tf.name

        with open(img_out_path, "ab") as fh:
            total = len(chunks)
            for idx, ch in enumerate(chunks):
                asyncio.run(
                    stream_chunk_to_file(
                        ch,
                        selected_voice,
                        rate_str,
                        volume_str,
                        fh,
                        img_progress,
                        idx,
                        total,
                    )
                )

        img_audio_bytes = open(img_out_path, "rb").read()
        st.audio(img_audio_bytes, format="audio/mp3")
        st.download_button(
            "💾 Download Images Audio MP3",
            data=img_audio_bytes,
            file_name="images_mastershub2023.mp3",
            mime="audio/mpeg",
        )
        img_progress.empty()
        st.success("Finished reading text from images.")

# ─────────────────────────────────────
# Read aloud with progress (manual text)
# ─────────────────────────────────────
st.markdown("---")
st.header("🔉 Read Aloud")

custom_text = st.text_area(
    "Paste or edit the text to read:",
    value=text_to_read,
    height=200,
    placeholder="Type or paste any text here to read aloud…",
)

c1, c2, c3 = st.columns([0.3, 0.3, 0.4])
with c1:
    start_btn = st.button("▶️ Read with Selected Voice")
with c2:
    dl_placeholder = st.empty()
with c3:
    st.caption("Progress is estimated by text chunks.")

if start_btn:
    if not custom_text.strip():
        st.warning("Please enter some text to read.")
    else:
        chunks = chunk_text(custom_text, max_len=700)
        progress = st.progress(0.0, text="Synthesizing…")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tf:
            out_path = tf.name

        with open(out_path, "ab") as fh:
            total = len(chunks)
            for idx, ch in enumerate(chunks):
                asyncio.run(
                    stream_chunk_to_file(
                        ch,
                        selected_voice,
                        rate_str,
                        volume_str,
                        fh,
                        progress,
                        idx,
                        total,
                    )
                )

        audio_bytes = open(out_path, "rb").read()
        st.audio(audio_bytes, format="audio/mp3")
        dl_placeholder.download_button(
            "💾 Download MP3",
            data=audio_bytes,
            file_name="speech_mastershub2023.mp3",
            mime="audio/mpeg",
        )
        progress.empty()
        st.success("Done!")

st.markdown("---")
st.caption(
    "© MastersHub2023 • Built with Streamlit, PyMuPDF, python-docx, pandas, edge-tts, "
    "and optional Tesseract OCR / Pillow."
)
